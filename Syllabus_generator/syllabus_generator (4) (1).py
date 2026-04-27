#!/usr/bin/env python3
"""
AI-Based Syllabus Generator
Powered by llama3.1:8b via Ollama
NBA Accreditation | OBE-Based Curriculum Format
"""

import json
import http.client
import os
import sys
import time
import threading

# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "127.0.0.1")
OLLAMA_PORT = int(os.environ.get("OLLAMA_PORT", "11434"))
MODEL       = "llama3.1:8b"

# ─────────────────────────────────────────────────────────────
# SPINNER
# ─────────────────────────────────────────────────────────────
class Spinner:
    def __init__(self, msg):
        self.msg     = msg
        self.running = False
        self.thread  = None

    def _spin(self):
        frames = ["⠋","⠙","⠹","⠸","⠼","⠴","⠦","⠧","⠇","⠏"]
        i = 0
        while self.running:
            print(f"\r{frames[i % len(frames)]}  {self.msg}", end="", flush=True)
            i += 1
            time.sleep(0.1)

    def start(self):
        self.running = True
        self.thread  = threading.Thread(target=self._spin, daemon=True)
        self.thread.start()

    def stop(self):
        self.running = False
        if self.thread:
            self.thread.join()
        print(f"\r✓  {self.msg}          ", flush=True)

# ─────────────────────────────────────────────────────────────
# OLLAMA CALL
# ─────────────────────────────────────────────────────────────
def call_ollama(prompt: str) -> str:
    payload = json.dumps({
        "model":  MODEL,
        "prompt": prompt,
        "stream": False
    }).encode("utf-8")

    try:
        conn = http.client.HTTPConnection(OLLAMA_HOST, OLLAMA_PORT, timeout=300)
        conn.request(
            "POST", "/api/generate",
            body=payload,
            headers={"Content-Type": "application/json",
                     "Content-Length": str(len(payload))}
        )
        res  = conn.getresponse()
        raw  = res.read().decode("utf-8")
        conn.close()

        # Ollama may send newline-delimited JSON chunks
        text = ""
        for line in raw.strip().splitlines():
            try:
                obj = json.loads(line)
                text += obj.get("response", "")
            except json.JSONDecodeError:
                pass

        return text.strip()

    except ConnectionRefusedError:
        print("\n\n[ERROR] Ollama is not running or not reachable.")
        print(f"  Host: {OLLAMA_HOST}:{OLLAMA_PORT}")
        print("  → Open a NEW terminal and run:  ollama serve")
        print("  → Then re-run this script.\n")
        sys.exit(1)
    except Exception as e:
        print(f"\n\n[ERROR] Ollama request failed: {e}\n")
        sys.exit(1)

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────
def divider(char="─", width=62):
    print(char * width)

def confirm(section: str) -> bool:
    while True:
        ans = input(f"\nAre you satisfied with {section}? (yes/no): ").strip().lower()
        if ans in ("yes", "y"):
            return True
        if ans in ("no", "n"):
            return False
        print("Please type yes or no.")

def generate_section(label: str, prompt: str) -> str:
    while True:
        sp = Spinner(f"Generating {label}...")
        sp.start()
        result = call_ollama(prompt)
        sp.stop()

        print()
        divider()
        print(result)
        divider()

        if confirm(label):
            return result
        print(f"\n[Regenerating {label}...]\n")

# ─────────────────────────────────────────────────────────────
# GENERATION PROMPTS
# ─────────────────────────────────────────────────────────────
def gen_objectives(info):
    prompt = f"""You are an academic syllabus expert for NBA-accredited engineering colleges.

Generate exactly 5 professional course objectives for:
Programme Level : {info['level']}
Programme       : {info['programme']}
Course Name     : {info['course']}

Rules:
- Use active academic verbs (To understand, To analyze, To implement…)
- Each objective must be one clear sentence
- Aligned with NBA/AICTE curriculum norms

Output format (copy exactly):
Course Objectives:
1. <objective>
2. <objective>
3. <objective>
4. <objective>
5. <objective>

Only output the list above. No extra text."""
    return generate_section("Course Objectives", prompt)


def gen_unit_titles(info):
    prompt = f"""You are an academic curriculum designer.

Generate exactly 5 progressive unit titles for:
Course : {info['course']}
Programme : {info['programme']}

Units must go from basic → advanced.

Output format (copy exactly, nothing else):
UNIT – I: <Fundamentals/Introduction title>
UNIT – II: <Core Concepts title>
UNIT – III: <Advanced Concepts title>
UNIT – IV: <Applications / Case Studies title>
UNIT – V: <Emerging Trends / Innovation title>

Only 5 lines. No extra text."""
    return generate_section("Unit Titles", prompt)


def gen_unit(info, num: int, title: str) -> str:
    roman = ["I", "II", "III", "IV", "V"]
    r = roman[num - 1] if num <= 5 else str(num)
    prompt = f"""You are an academic curriculum designer for Indian engineering universities (AICTE/NBA).

Generate the detailed syllabus content for the following unit exactly in the format used by Indian university syllabi (like VTU, AKTU, Anna University):

Course   : {info['course']}
Programme: {info['programme']}
Unit     : UNIT – {r}: {title}

STRICTLY follow this output format:

UNIT – {r}: {title}:  <subtopic 1>, <subtopic 2>, <subtopic 3>, <subtopic 4>, <subtopic 5>, <subtopic 6>, <subtopic 7>, <subtopic 8>, <subtopic 9>, <subtopic 10>.

Rules:
- Output ONLY one paragraph of comma-separated subtopics after the unit heading
- Include 8 to 12 specific, relevant subtopics/concepts
- Each subtopic is a short noun phrase (2–6 words), NOT a full sentence
- Do NOT use Bloom's taxonomy headings (no Remember/Understand/Apply etc.)
- Do NOT use bullet points, dashes, or numbered lists
- Do NOT add any explanation or extra text
- End with a full stop

Example format (for reference only):
UNIT – I: Introduction to Cyber Security:  Basic Cyber Security Concepts, Layers of Security, Vulnerability, Threat, Harmful Acts, Motive of Attackers, Active Attacks, Passive Attacks, Software Attacks, Hardware Attacks, Spectrum of Attacks, Taxonomy of Various Attacks, IP Spoofing, Methods of Defense, Security Models, Risk Management, CIA Triad.

Now generate for: UNIT – {r}: {title}"""
    return generate_section(f"Unit {num}: {title}", prompt)


def gen_outcomes(info, unit_titles: str) -> str:
    prompt = f"""You are an OBE (Outcome-Based Education) curriculum expert.

Generate exactly 5 measurable Course Outcomes for:
Course    : {info['course']}
Programme : {info['programme']}
Level     : {info['level']}

Units covered:
{unit_titles}

Rules:
- Each CO starts with "CO" + number
- Use a Bloom's Higher Order Thinking verb (Apply, Analyze, Evaluate, Design…)
- Must be measurable and assessable
- Aligned with NBA Programme Outcomes (POs)

Output format (copy exactly):
Course Outcomes:
CO1: <verb> <specific measurable outcome>
CO2: <verb> <specific measurable outcome>
CO3: <verb> <specific measurable outcome>
CO4: <verb> <specific measurable outcome>
CO5: <verb> <specific measurable outcome>

Only the list. No extra text."""
    return generate_section("Course Outcomes", prompt)


def gen_textbooks(info) -> str:
    prompt = f"""List exactly 3 well-known standard textbooks for:
{info['course']} | {info['programme']}

Output format:
Suggested Textbooks:
1. Author(s), "Book Title", Publisher, Edition, Year
2. Author(s), "Book Title", Publisher, Edition, Year
3. Author(s), "Book Title", Publisher, Edition, Year

Use only real published books. No extra text."""
    sp = Spinner("Generating Suggested Textbooks...")
    sp.start()
    result = call_ollama(prompt)
    sp.stop()
    print("\n" + result + "\n")
    return result


def gen_ref_books(info) -> str:
    prompt = f"""List exactly 3 reference books for:
{info['course']} | {info['programme']}

Output format:
Suggested Reference Books:
1. Author(s), "Book Title", Publisher, Edition, Year
2. Author(s), "Book Title", Publisher, Edition, Year
3. Author(s), "Book Title", Publisher, Edition, Year

Use only real published books. No extra text."""
    sp = Spinner("Generating Reference Books...")
    sp.start()
    result = call_ollama(prompt)
    sp.stop()
    print("\n" + result + "\n")
    return result


def gen_youtube(info) -> str:
    prompt = f"""Suggest 4 YouTube channels or NPTEL playlists for learning:
{info['course']}

Output format:
Relevant YouTube Resources:
1. <Channel/Playlist Name> - <Brief description> - <URL>
2. <Channel/Playlist Name> - <Brief description> - <URL>
3. <Channel/Playlist Name> - <Brief description> - <URL>
4. <Channel/Playlist Name> - <Brief description> - <URL>

Use real channels like NPTEL, MIT OpenCourseWare, freeCodeCamp, Neso Academy, etc.
Only the list. No extra text."""
    sp = Spinner("Generating YouTube Resources...")
    sp.start()
    result = call_ollama(prompt)
    sp.stop()
    print("\n" + result + "\n")
    return result

# ─────────────────────────────────────────────────────────────
# SAVE DOCX
# ─────────────────────────────────────────────────────────────
def save_docx(info, sections):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
    except ImportError:
        print("\n[ERROR] python-docx not installed.")
        print("  Run: pip install python-docx\n")
        sys.exit(1)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helper: set paragraph font ─────────────────────────────
    def set_run(run, size=11, bold=False, color=None, italic=False):
        run.font.name   = "Arial"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_heading(text, level=1, color=(31, 78, 121)):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        set_run(run, size=14 if level == 1 else 12, bold=True, color=color)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        return p

    def add_body(text, bold=False, color=None, size=11):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def add_bullet(text):
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        set_run(run, size=11)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def add_numbered(text):
        p   = doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        set_run(run, size=11)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "B0C4DE")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Title ──────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ACADEMIC SYLLABUS")
    set_run(r, size=18, bold=True, color=(31, 78, 121))
    t.paragraph_format.space_after = Pt(4)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("NBA Accreditation  |  OBE-Based Curriculum")
    set_run(r2, size=11, italic=True, color=(100, 100, 100))
    t2.paragraph_format.space_after = Pt(12)

    add_divider()

    # ── Info table ─────────────────────────────────────────────
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    labels = ["Programme Level", "Programme", "Course Name"]
    values = [info["level"], info["programme"], info["course"]]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = table.rows[i]
        # label cell
        lc = row.cells[0]
        lc.width = Inches(2)
        lp = lc.paragraphs[0]
        lr = lp.add_run(lbl)
        set_run(lr, bold=True, color=(31, 78, 121))
        # shading
        tc_pr = lc._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "DCE6F1")
        tc_pr.append(shd)
        # value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        set_run(vr)

    doc.add_paragraph()
    add_divider()

    # ── Course Objectives ──────────────────────────────────────
    add_heading("Course Objectives")
    for line in sections["objectives"].splitlines():
        line = line.strip()
        if not line or "Course Objectives" in line:
            continue
        clean = line.lstrip("•-*1234567890. ").strip()
        if clean:
            add_bullet(clean)

    add_divider()

    # ── Units ──────────────────────────────────────────────────
    for unit_text in sections["units"]:
        lines = [l.strip() for l in unit_text.splitlines() if l.strip()]
        if not lines:
            continue

        # First non-empty line = unit heading (e.g. "UNIT – I: Introduction to Cyber Security:")
        unit_heading_line = lines[0]

        # Split heading into "UNIT – I: Title" and the subtopics paragraph
        # The model may put heading and subtopics on same line or next line
        full_text = " ".join(lines)

        # Try to split at first colon-space after the title portion
        # Format: "UNIT – I: Title:  subtopic1, subtopic2, ..."
        import re as _re
        # Match: UNIT – <roman>: <Title>: <subtopics>
        m = _re.match(r"(UNIT\s*[-–]\s*[IVX]+\s*:\s*[^:]+):\s*(.*)", full_text, _re.DOTALL)
        if m:
            heading_part  = m.group(1).strip().rstrip(":")
            subtopic_part = m.group(2).strip()
        else:
            # Fallback: first line is heading, rest is content
            heading_part  = unit_heading_line
            subtopic_part = " ".join(lines[1:]).strip()

        # Render unit heading in bold blue
        p   = doc.add_paragraph()
        run = p.add_run(heading_part)
        set_run(run, size=12, bold=True, color=(31, 78, 121))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)

        # Render subtopics as a flowing paragraph (comma-separated)
        if subtopic_part:
            p2   = doc.add_paragraph()
            run2 = p2.add_run(subtopic_part)
            set_run(run2, size=11)
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after  = Pt(6)
            # Justify text like a real syllabus
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_divider()

    # ── Course Outcomes ────────────────────────────────────────
    add_heading("Course Outcomes")
    for line in sections["outcomes"].splitlines():
        line = line.strip()
        if not line or "Course Outcomes" in line:
            continue
        clean = line.lstrip("•-* ").strip()
        if clean:
            add_bullet(clean)

    add_divider()

    # ── Textbooks ──────────────────────────────────────────────
    add_heading("Suggested Textbooks")
    for line in sections["textbooks"].splitlines():
        line = line.strip()
        if not line or "Textbook" in line:
            continue
        clean = re.sub(r"^\d+\.\s*", "", line).strip()
        if clean:
            add_numbered(clean)

    add_divider()

    # ── Reference Books ────────────────────────────────────────
    add_heading("Suggested Reference Books")
    for line in sections["ref_books"].splitlines():
        line = line.strip()
        if not line or "Reference" in line:
            continue
        clean = re.sub(r"^\d+\.\s*", "", line).strip()
        if clean:
            add_numbered(clean)

    add_divider()

    # ── YouTube ────────────────────────────────────────────────
    add_heading("Relevant YouTube Resources")
    for line in sections["youtube"].splitlines():
        line = line.strip()
        if not line or "YouTube" in line:
            continue
        clean = re.sub(r"^\d+\.\s*", "", line).strip()
        if clean:
            add_numbered(clean)

    # ── Save ───────────────────────────────────────────────────
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in info["course"])
    filename  = f"{safe_name.replace(' ', '_')}_syllabus.docx"
    doc.save(filename)
    return os.path.abspath(filename)

# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────
def main():
    os.system("cls" if os.name == "nt" else "clear")

    print("═" * 62)
    print("  AI-Based Syllabus Generator")
    print("  Powered by llama3.1:8b via Ollama")
    print("  NBA Accreditation | OBE-Based Curriculum")
    print("═" * 62)
    print()

    level      = input("Enter Programme Level (e.g. Undergraduate)      : ").strip()
    programme  = input("Enter Programme       (e.g. B.Tech CSE)          : ").strip()
    course     = input("Enter Course Name     (e.g. Artificial Intelligence): ").strip()
    info       = {"level": level, "programme": programme, "course": course}

    print()
    print("═" * 62)
    print("  Starting Syllabus Generation...")
    print("═" * 62)
    print()

    # Step 1 – Objectives
    objectives = gen_objectives(info)

    # Step 2 – Unit Titles
    unit_titles_raw = gen_unit_titles(info)

    # Parse unit title lines
    unit_lines = [
        l.strip() for l in unit_titles_raw.splitlines()
        if l.strip().upper().startswith("UNIT")
    ]
    if not unit_lines:
        # fallback: take any 5 non-empty lines
        unit_lines = [l.strip() for l in unit_titles_raw.splitlines() if l.strip()][:5]

    # Step 3 – One unit at a time
    units = []
    for i, title in enumerate(unit_lines, 1):
        units.append(gen_unit(info, i, title))

    # Step 4 – Course Outcomes
    outcomes = gen_outcomes(info, "\n".join(unit_lines))

    # Step 5-7 – Auto sections
    print()
    divider("─")
    print("  Generating supporting sections (auto)...")
    divider("─")
    print()
    textbooks = gen_textbooks(info)
    ref_books = gen_ref_books(info)
    youtube   = gen_youtube(info)

    # Step 8-9 – Save DOCX
    print()
    print("═" * 62)
    print("  Compiling and saving Word document...")
    print("═" * 62)

    sp = Spinner("Saving syllabus to .docx...")
    sp.start()
    sections = {
        "objectives": objectives,
        "units":      units,
        "outcomes":   outcomes,
        "textbooks":  textbooks,
        "ref_books":  ref_books,
        "youtube":    youtube,
    }
    out_path = save_docx(info, sections)
    sp.stop()

    print()
    print("═" * 62)
    print("  [Syllabus generation completed successfully]")
    print(f"  [File saved: {out_path}]")
    print("═" * 62)
    print()


if __name__ == "__main__":
    main()
